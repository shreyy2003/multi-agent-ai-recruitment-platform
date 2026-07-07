import io
import logging
import PyPDF2
from docx import Document

logger = logging.getLogger(__name__)


# ── PDF ────────────────────────────────────────────────────────────────────────

def extract_text_from_pdf_bytes(file_bytes: bytes) -> str:
    """
    Extract text from a PDF supplied as raw bytes.
    No temp file written to disk.
    Falls back gracefully on pages that fail to extract.
    """
    text_parts = []

    try:
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))

        if reader.is_encrypted:
            raise ValueError(
                "PDF is password-protected and cannot be processed."
            )

        for i, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text.strip())
            except Exception as e:
                logger.warning(f"Could not extract text from PDF page {i + 1}: {e}")
                continue

    except ValueError:
        raise

    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        raise ValueError(f"Could not read PDF file: {e}")

    full_text = "\n".join(text_parts).strip()

    if not full_text:
        raise ValueError(
            "PDF appears to be scanned or image-based. "
            "No extractable text was found. Please paste the JD as text instead."
        )

    return full_text


# ── DOCX ───────────────────────────────────────────────────────────────────────

def extract_text_from_docx_bytes(file_bytes: bytes) -> str:
    """
    Extract text from a DOCX supplied as raw bytes.
    Captures paragraphs + table cells (tables are skipped by paragraph-only parsers).
    No temp file written to disk.
    """
    try:
        doc = Document(io.BytesIO(file_bytes))

        parts = []

        # Paragraphs
        for para in doc.paragraphs:
            stripped = para.text.strip()
            if stripped:
                parts.append(stripped)

        # Table cells — commonly missed by naive parsers
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                )
                if row_text:
                    parts.append(row_text)

    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        raise ValueError(f"Could not read DOCX file: {e}")

    full_text = "\n".join(parts).strip()

    if not full_text:
        raise ValueError(
            "DOCX file appears to be empty or contains no readable text."
        )

    return full_text