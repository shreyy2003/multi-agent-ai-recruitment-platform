import logging
import tempfile

from fastapi import APIRouter, UploadFile, File, Form, HTTPException, Body
from fastapi.responses import FileResponse

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    HRFlowable,
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors

from models.debrief_schemas import DebriefResponse
from agents.debrief_agent import generate_debrief
from utils.file_parser import (
    extract_text_from_pdf_bytes,
    extract_text_from_docx_bytes,
)

logger = logging.getLogger(__name__)

router = APIRouter(
    prefix="/agent",
    tags=["Debrief Agent"]
)


# ── Shared File Parser ─────────────────────────────────────────────────────────

def parse_uploaded_file(
    file_bytes: bytes,
    filename: str,
    label: str
) -> str:
    fname = filename.lower()
    try:
        if fname.endswith(".pdf"):
            return extract_text_from_pdf_bytes(file_bytes)
        elif fname.endswith(".docx"):
            return extract_text_from_docx_bytes(file_bytes)
        else:
            raise HTTPException(
                status_code=400,
                detail=f"{label}: only PDF and DOCX files are supported."
            )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"{label} could not be read: {str(e)}"
        )


# ── Main Debrief Endpoint ──────────────────────────────────────────────────────

@router.post("/debrief", response_model=DebriefResponse)
async def debrief_agent(
    candidate_name: str = Form(...),
    job_title: str = Form(...),
    interviewer_name: str = Form(...),
    jd_text: str = Form(""),
    jd_file: UploadFile | None = File(None),
    transcript_file: UploadFile = File(...),
    resume_file: UploadFile | None = File(None),
):
    # Resolve JD
    resolved_jd = jd_text.strip()
    if jd_file:
        file_bytes = await jd_file.read()
        resolved_jd = parse_uploaded_file(
            file_bytes, jd_file.filename or "", "JD file"
        )
    if not resolved_jd:
        raise HTTPException(
            status_code=400,
            detail="Provide either a JD text or upload a JD file."
        )

    # Resolve Transcript
    transcript_bytes = await transcript_file.read()
    resolved_transcript = parse_uploaded_file(
        transcript_bytes,
        transcript_file.filename or "",
        "Transcript file"
    )

    # Resolve Resume (optional)
    resolved_resume = None
    if resume_file:
        resume_bytes = await resume_file.read()
        resolved_resume = parse_uploaded_file(
            resume_bytes,
            resume_file.filename or "",
            "Resume file"
        )

    try:
        result = generate_debrief(
            candidate_name=candidate_name,
            job_title=job_title,
            interviewer_name=interviewer_name,
            jd_text=resolved_jd,
            transcript_text=resolved_transcript,
            resume_text=resolved_resume,
        )
    except ValueError as e:
        raise HTTPException(status_code=500, detail=str(e))

    return result


# ── DOCX Export ────────────────────────────────────────────────────────────────

@router.post("/debrief/export/docx")
async def export_debrief_docx(data: dict = Body(...)):
    try:
        r = data["report_data"]
        doc = Document()

        # Title
        title = doc.add_heading("Post-Interview Debrief Report", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.LEFT

        meta = doc.add_paragraph()
        meta.add_run("Candidate: ").bold = True
        meta.add_run(f"{r['candidate_name']}   ")
        meta.add_run("Role: ").bold = True
        meta.add_run(f"{r['job_title']}   ")
        meta.add_run("Interviewer: ").bold = True
        meta.add_run(r['interviewer_name'])
        doc.add_paragraph()

        # Executive Summary
        doc.add_heading("Executive Summary", level=2)
        doc.add_paragraph(r["executive_summary"])
        doc.add_paragraph()

        # JD Match
        doc.add_heading("JD Match Analysis", level=2)
        match_para = doc.add_paragraph()
        match_para.add_run("Overall Match: ").bold = True
        match_para.add_run(f"{r['jd_match']['match_percentage']}%")
        doc.add_paragraph()

        jd = r["jd_match"]

        if jd.get("matched_requirements"):
            doc.add_heading("Fully Met Requirements", level=3)
            for item in jd["matched_requirements"]:
                doc.add_paragraph(item, style="List Bullet")

        if jd.get("partial_matches"):
            doc.add_heading("Partially Met", level=3)
            for item in jd["partial_matches"]:
                doc.add_paragraph(item, style="List Bullet")

        if jd.get("missing_requirements"):
            doc.add_heading("Missing Requirements", level=3)
            for item in jd["missing_requirements"]:
                doc.add_paragraph(item, style="List Bullet")

        doc.add_paragraph()

        # Strengths
        doc.add_heading("Strengths", level=2)
        for item in r["strengths"]:
            doc.add_paragraph(item, style="List Bullet")
        doc.add_paragraph()

        # Concerns
        doc.add_heading("Concerns / Risks", level=2)
        for item in r["concerns"]:
            doc.add_paragraph(item, style="List Bullet")
        doc.add_paragraph()

        # Resume Consistency (optional — 3 sub-fields)
        if r.get("resume_uploaded") and r.get("resume_consistency"):
            rc = r["resume_consistency"]
            doc.add_heading("Resume Consistency Check", level=2)

            if rc.get("confirmed_skills"):
                doc.add_heading(
                    "Confirmed in Both Resume & Interview", level=3
                )
                for item in rc["confirmed_skills"]:
                    doc.add_paragraph(item, style="List Bullet")

            if rc.get("interview_claims_missing_from_resume"):
                doc.add_heading(
                    "Claimed in Interview — Not on Resume", level=3
                )
                for item in rc["interview_claims_missing_from_resume"]:
                    doc.add_paragraph(item, style="List Bullet")

            if rc.get("resume_skills_not_discussed"):
                doc.add_heading(
                    "On Resume — Not Discussed in Interview", level=3
                )
                for item in rc["resume_skills_not_discussed"]:
                    doc.add_paragraph(item, style="List Bullet")

            doc.add_paragraph()

        # Recommendation
        doc.add_heading("Final Recommendation", level=2)
        rec_para = doc.add_paragraph()
        rec_run = rec_para.add_run(r["final_recommendation"])
        rec_run.bold = True
        rec_run.font.size = Pt(14)
        doc.add_paragraph(r["recommendation_justification"])
        doc.add_paragraph()

        # Follow-Up Questions
        doc.add_heading("Suggested Follow-Up Questions", level=2)
        for i, q in enumerate(r["follow_up_questions"], 1):
            doc.add_paragraph(f"{i}. {q}", style="List Number")

        doc.add_paragraph()
        footer = doc.add_paragraph("TJJ - The Jobs Jungle  ·  Confidential")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
        doc.save(tmp.name)

        filename = (
            f"debrief_{r['candidate_name'].replace(' ', '_')}_"
            f"{r['job_title'].replace(' ', '_')}.docx"
        )

        return FileResponse(
            tmp.name,
            filename=filename,
            media_type=(
                "application/vnd.openxmlformats-"
                "officedocument.wordprocessingml.document"
            ),
        )

    except Exception as e:
        logger.error(f"Debrief DOCX export failed: {e}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail="Failed to generate DOCX report."
        )


# ── PDF Export ─────────────────────────────────────────────────────────────────

@router.post("/debrief/export/pdf")
async def export_debrief_pdf(data: dict = Body(...)):
    try:
        r = data["report_data"]

        tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")

        doc = SimpleDocTemplate(
            tmp.name,
            pagesize=A4,
            rightMargin=40,
            leftMargin=40,
            topMargin=40,
            bottomMargin=40,
        )

        styles = getSampleStyleSheet()
        VIOLET = colors.HexColor("#7C3AED")
        GRAY   = colors.HexColor("#6B7280")
        DARK   = colors.HexColor("#1E1E32")
        GREEN  = colors.HexColor("#34D399")
        YELLOW = colors.HexColor("#FBBF24")
        RED    = colors.HexColor("#F87171")

        title_style = ParagraphStyle(
            "DebriefTitle", parent=styles["Title"],
            textColor=VIOLET, fontSize=22, spaceAfter=6,
        )
        h1_style = ParagraphStyle(
            "DebriefH1", parent=styles["Heading1"],
            textColor=VIOLET, fontSize=13,
            spaceBefore=16, spaceAfter=6,
        )
        h2_style = ParagraphStyle(
            "DebriefH2", parent=styles["Heading2"],
            textColor=DARK, fontSize=10,
            spaceBefore=10, spaceAfter=4,
        )
        body_style = ParagraphStyle(
            "DebriefBody", parent=styles["Normal"],
            fontSize=9, leading=14, spaceAfter=4,
        )
        meta_style = ParagraphStyle(
            "DebriefMeta", parent=styles["Normal"],
            fontSize=9, textColor=GRAY, spaceAfter=8,
        )
        footer_style = ParagraphStyle(
            "DebriefFooter", parent=styles["Normal"],
            fontSize=8, textColor=GRAY, alignment=1,
        )

        def bullet(text, color=DARK):
            return Paragraph(
                f"<font color='#7C3AED'>›</font> {text}",
                ParagraphStyle(
                    "B", parent=styles["Normal"],
                    fontSize=9, leading=14,
                    leftIndent=12, spaceAfter=3,
                    textColor=color,
                )
            )

        elements = []

        # Title
        elements.append(
            Paragraph("Post-Interview Debrief Report", title_style)
        )
        elements.append(Paragraph(
            f"<b>Candidate:</b> {r['candidate_name']}  &nbsp;|&nbsp;  "
            f"<b>Role:</b> {r['job_title']}  &nbsp;|&nbsp;  "
            f"<b>Interviewer:</b> {r['interviewer_name']}",
            meta_style,
        ))
        elements.append(
            HRFlowable(width="100%", color=VIOLET, thickness=1)
        )
        elements.append(Spacer(1, 12))

        # Executive Summary
        elements.append(Paragraph("Executive Summary", h1_style))
        elements.append(Paragraph(r["executive_summary"], body_style))
        elements.append(Spacer(1, 10))

        # JD Match
        jd = r["jd_match"]
        elements.append(Paragraph("JD Match Analysis", h1_style))
        elements.append(Paragraph(
            f"<b>Overall JD Match: "
            f"<font color='#7C3AED'>{jd['match_percentage']}%</font></b>",
            body_style,
        ))
        elements.append(Spacer(1, 6))

        if jd.get("matched_requirements"):
            elements.append(
                Paragraph("Fully Met Requirements", h2_style)
            )
            for item in jd["matched_requirements"]:
                elements.append(bullet(item, GREEN))

        if jd.get("partial_matches"):
            elements.append(Paragraph("Partially Met", h2_style))
            for item in jd["partial_matches"]:
                elements.append(bullet(item, YELLOW))

        if jd.get("missing_requirements"):
            elements.append(
                Paragraph("Missing Requirements", h2_style)
            )
            for item in jd["missing_requirements"]:
                elements.append(bullet(item, RED))

        elements.append(Spacer(1, 10))

        # Strengths
        elements.append(Paragraph("Strengths", h1_style))
        for item in r["strengths"]:
            elements.append(bullet(item, GREEN))
        elements.append(Spacer(1, 10))

        # Concerns
        elements.append(Paragraph("Concerns / Risks", h1_style))
        for item in r["concerns"]:
            elements.append(bullet(item, YELLOW))
        elements.append(Spacer(1, 10))

        # Resume Consistency (optional — 3 sub-fields)
        if r.get("resume_uploaded") and r.get("resume_consistency"):
            rc = r["resume_consistency"]
            elements.append(
                Paragraph("Resume Consistency Check", h1_style)
            )

            if rc.get("confirmed_skills"):
                elements.append(
                    Paragraph(
                        "Confirmed in Both Resume & Interview", h2_style
                    )
                )
                for item in rc["confirmed_skills"]:
                    elements.append(bullet(item, GREEN))

            if rc.get("interview_claims_missing_from_resume"):
                elements.append(
                    Paragraph(
                        "Claimed in Interview — Not on Resume", h2_style
                    )
                )
                for item in rc["interview_claims_missing_from_resume"]:
                    elements.append(bullet(item, YELLOW))

            if rc.get("resume_skills_not_discussed"):
                elements.append(
                    Paragraph(
                        "On Resume — Not Discussed in Interview", h2_style
                    )
                )
                for item in rc["resume_skills_not_discussed"]:
                    elements.append(
                        bullet(item, colors.HexColor("#9CA3AF"))
                    )

            elements.append(Spacer(1, 10))

        # Recommendation
        elements.append(Paragraph("Final Recommendation", h1_style))
        elements.append(Paragraph(
            f"<b><font size='13' color='#7C3AED'>"
            f"{r['final_recommendation']}</font></b>",
            body_style,
        ))
        elements.append(
            Paragraph(r["recommendation_justification"], body_style)
        )
        elements.append(Spacer(1, 10))

        # Follow-Up Questions
        elements.append(
            Paragraph("Suggested Follow-Up Questions", h1_style)
        )
        for i, q in enumerate(r["follow_up_questions"], 1):
            elements.append(bullet(f"{i}. {q}"))

        elements.append(Spacer(1, 20))
        elements.append(
            HRFlowable(width="100%", color=GRAY, thickness=0.5)
        )
        elements.append(
            Paragraph(
                "TJJ - The Jobs Jungle  ·  Confidential",
                footer_style,
            )
        )

        doc.build(elements)

        filename = (
            f"debrief_{r['candidate_name'].replace(' ', '_')}_"
            f"{r['job_title'].replace(' ', '_')}.pdf"
        )

        return FileResponse(
            tmp.name,
            filename=filename,
            media_type="application/pdf",
        )

    except Exception as e:
        logger.error(f"Debrief PDF export failed: {e}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail="Failed to generate PDF report."
        )