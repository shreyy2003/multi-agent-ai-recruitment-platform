import logging
import tempfile
import json
import uuid

from fastapi import (
    APIRouter,
    UploadFile,
    File,
    HTTPException,
    Form,
    Body
)

from fastapi.responses import FileResponse

from docx import Document

from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer
)

from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import letter

from models.intake_schemas import IntakeResponse

from agents.intake_agent import (
    generate_intake_documents
)

from utils.file_parser import (
    extract_text_from_pdf_bytes,
    extract_text_from_docx_bytes
)

from database.db import get_db
from models.pipeline_schemas import JobRecord


# ── Router ──────────────────────────────────────────

router = APIRouter()


# ── Logger ──────────────────────────────────────────

logger = logging.getLogger(__name__)


# ── Config ──────────────────────────────────────────

MAX_FILE_SIZE_MB = 5

SUPPORTED_EXTS = {
    "pdf",
    "docx"
}


# ────────────────────────────────────────────────────
# AGENT 1 — INTAKE AGENT
# ────────────────────────────────────────────────────

@router.post(
    "/agent/intake",
    response_model=IntakeResponse
)
async def intake_agent(
    file: UploadFile = File(None),
    jd_text: str = Form(None)
):

    try:

        # ── Text Input ───────────────────────────────

        if jd_text and jd_text.strip():

            logger.info("Processing pasted JD text")

            extracted_jd_text = jd_text.strip()

        # ── File Upload ──────────────────────────────

        elif file and file.filename:

            logger.info(
                f"Processing uploaded file: {file.filename}"
            )

            ext = file.filename.rsplit(
                ".",
                1
            )[-1].lower()

            if ext not in SUPPORTED_EXTS:

                raise HTTPException(
                    status_code=400,
                    detail=(
                        f"Unsupported file type '.{ext}'. "
                        f"Only PDF and DOCX files are accepted."
                    )
                )

            # ── File Size Validation ────────────────

            file.file.seek(0, 2)

            size_mb = (
                file.file.tell() /
                (1024 * 1024)
            )

            file.file.seek(0)

            if size_mb > MAX_FILE_SIZE_MB:

                raise HTTPException(
                    status_code=413,
                    detail=(
                        f"File is too large "
                        f"({size_mb:.1f}MB). "
                        f"Maximum allowed size is "
                        f"{MAX_FILE_SIZE_MB}MB."
                    )
                )

            # ── Read File ───────────────────────────

            file_bytes = await file.read()

            try:

                if ext == "pdf":

                    extracted_jd_text = (
                        extract_text_from_pdf_bytes(
                            file_bytes
                        )
                    )

                else:

                    extracted_jd_text = (
                        extract_text_from_docx_bytes(
                            file_bytes
                        )
                    )

            except ValueError as e:

                raise HTTPException(
                    status_code=422,
                    detail=str(e)
                )

        # ── No Input ─────────────────────────────────

        else:

            raise HTTPException(
                status_code=400,
                detail=(
                    "Provide either JD text "
                    "or upload a PDF/DOCX file."
                )
            )

        # ── Empty Content Validation ────────────────

        if not extracted_jd_text.strip():

            raise HTTPException(
                status_code=422,
                detail=(
                    "The provided JD appears to be empty. "
                    "Please check your input and try again."
                )
            )

        logger.info(
            f"JD ready for processing "
            f"({len(extracted_jd_text)} characters)"
        )

        # ── Generate Intake Documents ───────────────

        result = await generate_intake_documents(
            extracted_jd_text
        )

        logger.info(
            f"Intake documents generated successfully "
            f"| Tokens: {result.tokens_used}"
        )

        return result

    except HTTPException:
        raise

    except Exception as e:

        logger.error(
            f"Unhandled server error: {e}",
            exc_info=True
        )

        raise HTTPException(
            status_code=500,
            detail=(
                "Something went wrong processing "
                "your request. "
                "Please try again later."
            )
        )


# ────────────────────────────────────────────────────
# SAVE TO PIPELINE — optional, called after intake result is shown
# ────────────────────────────────────────────────────

@router.post(
    "/agent/intake/save-to-pipeline",
    response_model=JobRecord
)
async def save_intake_to_pipeline(data: dict = Body(...)):
    """
    Saves an already-generated intake result as a job record.
    Frontend sends: { job_title, structured_jd (full intake result) }
    """

    job_title = data.get("job_title")
    structured_jd = data.get("structured_jd")

    if not job_title or not structured_jd:
        raise HTTPException(
            status_code=400,
            detail="job_title and structured_jd are required."
        )

    # Domain sent from frontend (recruiter selects from dropdown before saving)
    domain = data.get("domain") or None

    job_id = str(uuid.uuid4())[:12]

    try:

        with get_db() as conn:

            conn.execute(
                """
                INSERT INTO jobs (job_id, job_title, structured_jd, domain, status)
                VALUES (?, ?, ?, ?, 'open')
                """,
                (job_id, job_title, json.dumps(structured_jd), domain),
            )

            row = conn.execute(
                "SELECT * FROM jobs WHERE job_id = ?", (job_id,)
            ).fetchone()

        logger.info(f"Job saved to pipeline: {job_id} | {job_title}")

        return JobRecord(
            job_id=row["job_id"],
            job_title=row["job_title"],
            structured_jd=json.loads(row["structured_jd"]),
            domain=row["domain"],
            status=row["status"],
            created_at=row["created_at"],
            updated_at=row["updated_at"],
            candidate_count=0,
        )

    except Exception as e:

        logger.error(f"Failed to save job to pipeline: {e}", exc_info=True)

        raise HTTPException(
            status_code=500,
            detail="Failed to save job to pipeline."
        )


# ────────────────────────────────────────────────────
# DOWNLOAD REPORT
# ────────────────────────────────────────────────────

@router.post("/download-report")
async def download_report(
    data: dict = Body(...)
):

    doc = Document()

    # ── Title ──────────────────────────────────────

    doc.add_heading(
        "Recruitment Intake Report",
        level=1
    )

    # ── Hiring Context ─────────────────────────────

    hiring = data.get(
        "hiring_context",
        {}
    )

    doc.add_heading(
        "Hiring Context",
        level=2
    )

    for key, value in hiring.items():

        if isinstance(value, list):
            value = ", ".join(value)

        doc.add_paragraph(
            f"{key.replace('_', ' ').title()}: {value}"
        )

    # ── Intake Document ────────────────────────────

    intake = data.get(
        "intake_document",
        {}
    )

    doc.add_heading(
        "Intake Meeting Document",
        level=2
    )

    for key, value in intake.items():

        heading = key.replace(
            "_",
            " "
        ).title()

        doc.add_heading(
            heading,
            level=3
        )

        if isinstance(value, list):

            for item in value:

                doc.add_paragraph(
                    item,
                    style="List Bullet"
                )

        else:

            doc.add_paragraph(
                str(value)
            )

    # ── Auto Reject Criteria ───────────────────────

    auto_reject = data.get(
        "auto_reject_criteria",
        []
    )

    doc.add_heading(
        "Auto Reject Criteria",
        level=2
    )

    for item in auto_reject:

        doc.add_paragraph(
            item,
            style="List Bullet"
        )

    # ── Prescreening Questions ─────────────────────

    questions = data.get(
        "prescreening_questions",
        []
    )

    doc.add_heading(
        "Prescreening Questions",
        level=2
    )

    for idx, q in enumerate(
        questions,
        start=1
    ):

        doc.add_heading(
            f"Q{idx}: {q.get('question')}",
            level=3
        )

        doc.add_paragraph(
            f"Category: {q.get('category')}"
        )

        doc.add_paragraph(
            f"What To Listen For: "
            f"{q.get('listen_for')}"
        )

    # ── Scorecard ──────────────────────────────────

    scorecard = data.get(
        "scorecard",
        {}
    )

    doc.add_heading(
        "Candidate Scorecard",
        level=2
    )

    for item in scorecard.get(
        "criteria",
        []
    ):

        doc.add_heading(
            item.get("name"),
            level=3
        )

        doc.add_paragraph(
            f"Weight: {item.get('weight')}%"
        )

        doc.add_paragraph(
            item.get("description")
        )

        guide = item.get(
            "scoring_guide"
        )

        if guide:

            doc.add_paragraph(
                f"Score 5: "
                f"{guide.get('score_5')}"
            )

            doc.add_paragraph(
                f"Score 3: "
                f"{guide.get('score_3')}"
            )

            doc.add_paragraph(
                f"Score 1: "
                f"{guide.get('score_1')}"
            )

    doc.add_heading(
        "Recommendation Threshold",
        level=3
    )

    doc.add_paragraph(
        scorecard.get(
            "recommendation_threshold",
            ""
        )
    )

    # ── Save Temp File ─────────────────────────────

    temp_file = tempfile.NamedTemporaryFile(
        delete=False,
        suffix=".docx"
    )

    doc.save(temp_file.name)

    return FileResponse(
        temp_file.name,
        filename="intake_report.docx",
        media_type=(
            "application/vnd.openxmlformats-"
            "officedocument.wordprocessingml.document"
        )
    )


# ────────────────────────────────────────────────────
# EXPORT DOCX
# ────────────────────────────────────────────────────

@router.post("/agent/export/docx")
async def export_docx(data: dict):

    try:

        response = data["report_data"]

        doc = Document()

        doc.add_heading(
            "TJJ Recruitment Intake Report",
            level=1
        )

        hiring_context = response.get(
            "hiring_context",
            {}
        )

        doc.add_heading(
            "Hiring Context",
            level=2
        )

        for key, value in hiring_context.items():

            if isinstance(value, list):
                value = ", ".join(value)

            if value:

                doc.add_paragraph(
                    f"{key.replace('_', ' ').title()}: {value}"
                )

        intake = response.get(
            "intake_document",
            {}
        )

        doc.add_heading(
            "Intake Meeting Document",
            level=2
        )

        for section, content in intake.items():

            doc.add_heading(
                section.replace("_", " ").title(),
                level=3
            )

            if isinstance(content, list):

                for item in content:

                    doc.add_paragraph(
                        item,
                        style="List Bullet"
                    )

            else:

                doc.add_paragraph(
                    str(content)
                )

        temp_file = tempfile.NamedTemporaryFile(
            delete=False,
            suffix=".docx"
        )

        doc.save(temp_file.name)

        logger.info(
            "DOCX report generated successfully"
        )

        return FileResponse(
            temp_file.name,
            media_type=(
                "application/vnd.openxmlformats-"
                "officedocument.wordprocessingml.document"
            ),
            filename="intake_report.docx"
        )

    except Exception as e:

        logger.error(
            f"DOCX export failed: {e}",
            exc_info=True
        )

        raise HTTPException(
            status_code=500,
            detail="Failed to generate DOCX report."
        )


# ────────────────────────────────────────────────────
# EXPORT PDF
# ────────────────────────────────────────────────────

@router.post("/agent/export/pdf")
async def export_pdf(data: dict):

    try:

        response = data["report_data"]

        temp_file = tempfile.NamedTemporaryFile(
            delete=False,
            suffix=".pdf"
        )

        doc = SimpleDocTemplate(
            temp_file.name,
            pagesize=letter,
            rightMargin=40,
            leftMargin=40,
            topMargin=40,
            bottomMargin=40
        )

        styles = getSampleStyleSheet()

        elements = []

        elements.append(
            Paragraph(
                "TJJ Recruitment Intake Report",
                styles["Title"]
            )
        )

        elements.append(
            Spacer(1, 20)
        )

        doc.build(elements)

        logger.info(
            "PDF report generated successfully"
        )

        return FileResponse(
            temp_file.name,
            media_type="application/pdf",
            filename="intake_report.pdf"
        )

    except Exception as e:

        logger.error(
            f"PDF export failed: {e}",
            exc_info=True
        )

        raise HTTPException(
            status_code=500,
            detail="Failed to generate PDF report."
        )